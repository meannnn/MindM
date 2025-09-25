#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
AI文档助手 - 主应用文件
基于Flask的Web应用，集成阿里云百炼API和Word文档处理功能
"""

import os
import sys
import tempfile
import atexit
import time
from datetime import datetime
from flask import Flask, render_template, request, jsonify, send_file
from werkzeug.utils import secure_filename
from dotenv import load_dotenv

# 添加项目根目录到Python路径
sys.path.append(os.path.dirname(os.path.abspath(__file__)))

# 导入自定义模块
try:
    from llm.clients.aliyun_client import AliyunClient
    from llm.utils.file_handler import FileHandler
    from docx_processor.template_processor import TemplateProcessor
    from prompts.teaching_design_prompt import get_teaching_design_prompt
    from schemas.teaching_design_schema import validate_teaching_design_data
    from utils.logger import setup_logger, get_logger, DocumentProcessingLogger, timing_decorator
except ImportError as e:
    print(f"导入模块失败: {e}")
    print("请确保所有模块已正确安装")
    sys.exit(1)

# 加载环境变量
load_dotenv()

# 配置集中式日志系统
logger = setup_logger(
    name="ai_doc_assistant",
    log_level=os.getenv("LOG_LEVEL", "DEBUG"),  # 改为DEBUG以获得更详细日志
    log_file=os.getenv("LOG_FILE", "./logs/app.log"),
    console_output=True,
    enable_colors=True
)

# 创建Flask应用
app = Flask(__name__, 
           template_folder='web/templates',
           static_folder='web/static')
app.config['SECRET_KEY'] = os.getenv('SECRET_KEY', 'dev-secret-key')
app.config['MAX_CONTENT_LENGTH'] = 16 * 1024 * 1024  # 16MB max file size

# 初始化组件
try:
    llm_client = AliyunClient()
    file_handler = FileHandler()
    template_processor = TemplateProcessor()
    logger.info("组件初始化成功")
except Exception as e:
    logger.error(f"组件初始化失败: {e}")
    sys.exit(1)

# 存储文件信息的简单内存存储（生产环境应使用数据库）
uploaded_files = {}

# 存储临时文件路径，用于清理
temp_files = set()

def cleanup_temp_files():
    """清理临时文件"""
    for temp_path in temp_files.copy():
        try:
            if os.path.exists(temp_path):
                os.remove(temp_path)
                temp_files.discard(temp_path)
                logger.info(f"清理临时文件: {temp_path}")
        except Exception as e:
            logger.error(f"清理临时文件失败: {temp_path}, 错误: {e}")

# 注册清理函数
atexit.register(cleanup_temp_files)

@app.route('/')
def index():
    """主页"""
    return render_template('index.html')

@app.route('/chat')
def chat():
    """聊天界面"""
    return render_template('chat.html')

@app.route('/upload_file', methods=['POST'])
def upload_file():
    """文件上传接口"""
    try:
        if 'file' not in request.files:
            return jsonify({'success': False, 'error': '没有选择文件'})
        
        file = request.files['file']
        if file.filename == '':
            return jsonify({'success': False, 'error': '没有选择文件'})
        
        # 保存文件
        file_content = file.read()
        result = file_handler.save_uploaded_file(file_content, file.filename)
        
        if result['success']:
            file_id = result['file_id']
            uploaded_files[file_id] = result
            
            # 提取文本内容
            text_result = file_handler.extract_text_from_docx(result['file_path'])
            if text_result['success']:
                uploaded_files[file_id]['text_content'] = text_result['text_content']
            
            return jsonify({
                'success': True,
                'file_id': file_id,
                'filename': result['original_filename'],
                'file_size': result['file_size']
            })
        else:
            return jsonify({'success': False, 'error': result['message']})
            
    except Exception as e:
        logger.error(f"文件上传失败: {e}")
        return jsonify({'success': False, 'error': str(e)})

@app.route('/upload', methods=['POST'])
@timing_decorator("file_upload_and_processing")
def upload():
    """文件上传接口（兼容前端）"""
    import uuid
    task_id = str(uuid.uuid4())
    doc_logger = DocumentProcessingLogger(task_id, logger)
    
    try:
        doc_logger.logger.info(f"🚀 STARTING FILE UPLOAD PROCESS - Task ID: {task_id}")
        
        # 验证文件
        if 'file' not in request.files:
            doc_logger.logger.error("❌ NO FILE PROVIDED")
            return jsonify({'success': False, 'error': '没有选择文件'})
        
        file = request.files['file']
        if file.filename == '':
            doc_logger.logger.error("❌ EMPTY FILENAME")
            return jsonify({'success': False, 'error': '没有选择文件'})
        
        # 获取其他参数
        template = request.form.get('template', '思维发展型课堂教学设计模板')
        ai_model = request.form.get('ai_model', 'DeepSeek Chat')
        
        doc_logger.logger.info(f"📋 PROCESSING PARAMETERS - Template: {template}, AI Model: {ai_model}")
        
        # 保存文件
        file_content = file.read()
        doc_logger.log_upload_start(file.filename, len(file_content))
        
        result = file_handler.save_uploaded_file(file_content, file.filename)
        
        if result['success']:
            doc_logger.log_upload_complete(True, f"Saved as {result['saved_filename']}")
            
            file_id = result['file_id']
            uploaded_files[file_id] = result
            
            # 提取文本内容
            doc_logger.log_text_extraction_start()
            text_result = file_handler.extract_text_from_docx(result['file_path'])
            if text_result['success']:
                uploaded_files[file_id]['text_content'] = text_result['text_content']
                doc_logger.log_text_extraction_complete(True, len(text_result['text_content']))
            else:
                doc_logger.log_text_extraction_complete(False)
                doc_logger.logger.error(f"❌ TEXT EXTRACTION FAILED: {text_result.get('message', 'Unknown error')}")
            
            # 存储任务信息
            uploaded_files[file_id]['task_id'] = task_id
            uploaded_files[file_id]['template'] = template
            uploaded_files[file_id]['ai_model'] = ai_model
            uploaded_files[file_id]['status'] = 'processing'
            
            # 立即开始AI处理
            try:
                doc_logger.logger.info(f"🤖 STARTING AI PROCESSING - File ID: {file_id}")
                
                # 使用Qwen-Long模型和文件上传方式生成教学设计
                # 上传用户文件到阿里云
                doc_logger.logger.info("📤 UPLOADING USER FILE TO ALIYUN")
                user_upload_result = llm_client.upload_file_with_openai_client(uploaded_files[file_id]['file_path'])
                
                if not user_upload_result['success']:
                    doc_logger.logger.error(f"❌ USER FILE UPLOAD FAILED: {user_upload_result.get('message', 'Unknown error')}")
                    uploaded_files[file_id]['status'] = 'failed'
                    uploaded_files[file_id]['error'] = f'用户文件上传失败: {user_upload_result.get("message", "未知错误")}'
                    return jsonify({
                        'success': False,
                        'error': uploaded_files[file_id]['error']
                    })
                
                doc_logger.logger.info(f"✅ USER FILE UPLOADED SUCCESSFULLY - File ID: {user_upload_result['file_id']}")
                
                # 上传模板文件到阿里云
                doc_logger.logger.info("📤 UPLOADING TEMPLATE FILE TO ALIYUN")
                template_path = file_handler.get_template_file_content()['template_path']
                template_upload_result = llm_client.upload_file_with_openai_client(template_path)
                
                if not template_upload_result['success']:
                    doc_logger.logger.error(f"❌ TEMPLATE FILE UPLOAD FAILED: {template_upload_result.get('message', 'Unknown error')}")
                    uploaded_files[file_id]['status'] = 'failed'
                    uploaded_files[file_id]['error'] = f'模板文件上传失败: {template_upload_result.get("message", "未知错误")}'
                    return jsonify({
                        'success': False,
                        'error': uploaded_files[file_id]['error']
                    })
                
                doc_logger.logger.info(f"✅ TEMPLATE FILE UPLOADED SUCCESSFULLY - File ID: {template_upload_result['file_id']}")
                
                # 使用Qwen-Long模型和文件ID进行对话
                doc_logger.log_ai_processing_start("qwen-long")
                doc_logger.log_ai_processing_stage("file_preparation", "Files uploaded, starting AI conversation")
                
                ai_result = llm_client.chat_with_qwen_long_and_files(
                    user_file_id=user_upload_result['file_id'],
                    template_file_id=template_upload_result['file_id'],
                    model="qwen-long"
                )
                
                if ai_result['success']:
                    response_length = len(ai_result.get('text', ''))
                    doc_logger.log_ai_processing_complete(True, response_length, "qwen-long")
                    
                    doc_logger.logger.info("✅ AI PROCESSING SUCCESSFUL - Updating file status")
                    uploaded_files[file_id]['status'] = 'completed'
                    uploaded_files[file_id]['result'] = ai_result['text']
                    uploaded_files[file_id]['ai_response'] = ai_result
                    uploaded_files[file_id]['user_file_id'] = user_upload_result['file_id']
                    uploaded_files[file_id]['template_file_id'] = template_upload_result['file_id']
                    
                    doc_logger.logger.info(f"🎉 COMPLETE SUCCESS - Task ID: {task_id}, Response length: {response_length}")
                    
                    return jsonify({
                        'success': True,
                        'task_id': task_id,
                        'file_id': file_id,
                        'filename': result['original_filename'],
                        'file_size': result['file_size'],
                        'template': template,
                        'ai_model': ai_model,
                        'status': 'completed',
                        'result': ai_result['text'],
                        'request_id': ai_result.get('request_id'),
                        'usage': ai_result.get('usage'),
                        'model': ai_result.get('model')
                    })
                else:
                    doc_logger.log_ai_processing_complete(False, 0, "qwen-long")
                    doc_logger.logger.error(f"❌ AI PROCESSING FAILED: {ai_result.get('message', 'Unknown error')}")
                    uploaded_files[file_id]['status'] = 'failed'
                    uploaded_files[file_id]['error'] = ai_result.get('message', 'AI处理失败')
                    return jsonify({
                        'success': False,
                        'task_id': task_id,
                        'file_id': file_id,
                        'error': uploaded_files[file_id]['error']
                    })
                    
            except Exception as e:
                doc_logger.logger.error(f"💥 AI PROCESSING EXCEPTION: {str(e)}", exc_info=True)
                uploaded_files[file_id]['status'] = 'failed'
                uploaded_files[file_id]['error'] = str(e)
                return jsonify({
                    'success': False,
                    'task_id': task_id,
                    'file_id': file_id,
                    'error': str(e)
                })
        else:
            doc_logger.log_upload_complete(False, result.get('message', 'Unknown error'))
            return jsonify({'success': False, 'error': result['message']})
            
    except Exception as e:
        doc_logger.logger.error(f"💥 UPLOAD PROCESS EXCEPTION: {str(e)}", exc_info=True)
        return jsonify({'success': False, 'error': str(e)})

@app.route('/send_message', methods=['POST'])
def send_message():
    """发送消息接口"""
    try:
        data = request.get_json()
        message = data.get('message')
        file_id = data.get('file_id')
        conversation_history = data.get('history', [])
        
        if not message:
            return jsonify({'success': False, 'error': '消息不能为空'})
        
        # 获取文件信息
        file_info = None
        if file_id and file_id in uploaded_files:
            file_info = uploaded_files[file_id]
        
        # 检查是否是教学设计生成请求
        is_teaching_design_request = (
            file_info and 
            file_info.get('text_content') and 
            ('教学设计' in message or '教学' in message or '教案' in message)
        )
        
        if is_teaching_design_request:
            # 使用双文件模式生成教学设计
            logger.info("使用双文件模式生成教学设计")
            
            # 获取模板文件内容
            template_result = file_handler.get_template_file_content()
            
            if not template_result['success']:
                return jsonify({
                    'success': False,
                    'error': f'获取模板文件失败: {template_result.get("message", "未知错误")}'
                })
            
            # 调用双文件模式的AI模型
            result = llm_client.call_model_with_files(
                user_file_content=file_info['text_content'],
                template_file_content=template_result['template_content']
            )
            
            if result['success']:
                return jsonify({
                    'success': True,
                    'response': result['text'],
                    'request_id': result.get('request_id'),
                    'usage': result.get('usage'),
                    'mode': 'teaching_design'
                })
            else:
                return jsonify({
                    'success': False,
                    'error': result.get('message', 'AI处理失败'),
                    'status_code': result.get('status_code')
                })
        else:
            # 使用原有的单文件模式
            # 构建提示词
            prompt_content = message
            if file_info and file_info.get('text_content'):
                prompt_content = f"以下是文件内容：\n{file_info['text_content']}\n\n我的问题是：{message}"
            
            # 调用AI模型
            if conversation_history:
                # 构建完整的对话历史
                messages = []
                for msg in conversation_history:
                    messages.append({
                        "role": msg['role'],
                        "content": msg['content']
                    })
                # 添加当前消息
                messages.append({
                    "role": "user",
                    "content": prompt_content
                })
                result = llm_client.call_model_with_history(messages)
            else:
                result = llm_client.call_model(prompt_content)
            
            if result['success']:
                return jsonify({
                    'success': True,
                    'response': result['text'],
                    'request_id': result.get('request_id'),
                    'usage': result.get('usage'),
                    'mode': 'normal'
                })
            else:
                return jsonify({
                    'success': False,
                    'error': result.get('message', 'AI处理失败'),
                    'status_code': result.get('status_code')
                })
            
    except Exception as e:
        logger.error(f"发送消息失败: {e}")
        return jsonify({'success': False, 'error': str(e)})

@app.route('/get_file_content/<file_id>')
def get_file_content(file_id):
    """获取文件内容"""
    try:
        if file_id not in uploaded_files:
            return jsonify({'success': False, 'error': '文件不存在'})
        
        file_info = uploaded_files[file_id]
        return jsonify({
            'success': True,
            'filename': file_info['original_filename'],
            'text_content': file_info.get('text_content', ''),
            'file_size': file_info['file_size']
        })
        
    except Exception as e:
        logger.error(f"获取文件内容失败: {e}")
        return jsonify({'success': False, 'error': str(e)})

@app.route('/generate_teaching_design', methods=['POST'])
def generate_teaching_design():
    """生成教学设计"""
    try:
        data = request.get_json()
        file_id = data.get('file_id')
        template_id = data.get('template_id', 'default')
        
        if not file_id or file_id not in uploaded_files:
            return jsonify({'success': False, 'error': '文件不存在'})
        
        file_info = uploaded_files[file_id]
        
        if not file_info.get('text_content'):
            return jsonify({'success': False, 'error': '文件内容为空'})
        
        # 根据模板ID选择模板路径
        template_paths = {
            'default': 'templates/teaching_design_template.docx',
            'optimized': 'templates/optimized_teaching_design_template.docx',
            'table': 'templates/table_teaching_design_template.docx',
            'advanced_table': 'templates/advanced_table_teaching_design_template.docx',
            'docxtpl_table': 'templates/docxtpl_table_teaching_design_template.docx',
            'final_correct': 'templates/final_correct_template.docx',
            'final_table': 'templates/final_table_template.docx'
        }
        
        selected_template = template_paths.get(template_id, template_paths['default'])
        template_processor.set_template(selected_template)
        
        # 获取模板文件内容
        template_result = file_handler.get_template_file_content()
        if not template_result['success']:
            return jsonify({
                'success': False,
                'error': f'获取模板文件失败: {template_result.get("message", "未知错误")}'
            })
        
        # 构建提示词
        prompt = get_teaching_design_prompt(
            file_info['text_content'],
            template_result['template_content']
        )
        
        # 调用LLM生成教学设计
        llm_result = llm_client.call_model(prompt)
        
        if not llm_result['success']:
            return jsonify({
                'success': False,
                'error': f'LLM调用失败: {llm_result.get("message", "未知错误")}'
            })
        
        # 解析LLM返回的JSON
        try:
            import json
            json_data = llm_result['text']
            
            # 清理JSON数据（移除可能的markdown标记）
            if '```json' in json_data:
                json_data = json_data.split('```json')[1].split('```')[0].strip()
            elif '```' in json_data:
                json_data = json_data.split('```')[1].split('```')[0].strip()
            
            # 验证JSON格式
            is_valid, errors = validate_teaching_design_data(json.loads(json_data))
            if not is_valid:
                return jsonify({
                    'success': False,
                    'error': 'LLM返回的数据格式不正确',
                    'validation_errors': errors
                })
            
            # 生成Word文档
            output_filename = f"教学设计_{file_info['original_filename']}"
            output_path = os.path.join('outputs', output_filename)
            
            template_result = template_processor.process_teaching_design(json_data, output_path)
            
            if template_result['success']:
                # 存储生成结果
                file_info['generated_design'] = {
                    'json_data': json_data,
                    'output_path': template_result['output_path'],
                    'file_size': template_result['file_size'],
                    'generated_time': template_result['generated_time']
                }
                
                return jsonify({
                    'success': True,
                    'message': '教学设计生成成功',
                    'output_path': template_result['output_path'],
                    'file_size': template_result['file_size'],
                    'download_url': f'/download_design/{file_id}'
                })
            else:
                return jsonify({
                    'success': False,
                    'error': f'模板处理失败: {template_result.get("message", "未知错误")}'
                })
                
        except json.JSONDecodeError as e:
            return jsonify({
                'success': False,
                'error': f'JSON解析失败: {str(e)}',
                'raw_response': llm_result['text'][:500] + '...' if len(llm_result['text']) > 500 else llm_result['text']
            })
        except Exception as e:
            return jsonify({
                'success': False,
                'error': f'处理失败: {str(e)}'
            })
            
    except Exception as e:
        logger.error(f"生成教学设计失败: {e}")
        return jsonify({'success': False, 'error': str(e)})

@app.route('/download_design/<file_id>')
def download_design(file_id):
    """下载生成的教学设计文档"""
    try:
        if file_id not in uploaded_files:
            return jsonify({'success': False, 'error': '文件不存在'}), 404
        
        file_info = uploaded_files[file_id]
        
        if not file_info.get('generated_design'):
            return jsonify({'success': False, 'error': '教学设计尚未生成'}), 400
        
        design_info = file_info['generated_design']
        output_path = design_info['output_path']
        
        if not os.path.exists(output_path):
            return jsonify({'success': False, 'error': '生成的文件不存在'}), 404
        
        # 返回文件
        response = send_file(
            output_path,
            as_attachment=True,
            download_name=f"教学设计_{file_info['original_filename']}",
            mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
        )
        
        return response
        
    except Exception as e:
        logger.error(f"下载教学设计失败: {e}")
        return jsonify({'success': False, 'error': str(e)}), 500

@app.route('/templates')
def get_templates():
    """获取可用模板列表"""
    try:
        templates = [
            {
                'id': 'default',
                'name': '思维发展型课堂教学设计模板（基础版）',
                'description': '基于思维发展型课堂理念的教学设计模板，使用基础占位符',
                'path': 'templates/teaching_design_template.docx'
            },
            {
                'id': 'optimized',
                'name': '思维发展型课堂教学设计模板（优化版）',
                'description': '优化版教学设计模板，改进的格式和布局',
                'path': 'templates/optimized_teaching_design_template.docx'
            },
            {
                'id': 'table',
                'name': '思维发展型课堂教学设计模板（表插入版）',
                'description': '使用表插入形式展示学习活动的教学设计模板',
                'path': 'templates/table_teaching_design_template.docx'
            },
            {
                'id': 'advanced_table',
                'name': '思维发展型课堂教学设计模板（高级表插入版）',
                'description': '高级表插入形式的教学设计模板，支持复杂的表格循环',
                'path': 'templates/advanced_table_teaching_design_template.docx'
            },
            {
                'id': 'docxtpl_table',
                'name': '思维发展型课堂教学设计模板（docxtpl表插入版）',
                'description': '使用docxtpl表插入功能的教学设计模板，支持动态表格生成',
                'path': 'templates/docxtpl_table_teaching_design_template.docx'
            },
            {
                'id': 'final_correct',
                'name': '思维发展型课堂教学设计模板（修复版）',
                'description': '修复了表格插入问题的教学设计模板，支持正确的多条数据插入',
                'path': 'templates/final_correct_template.docx'
            },
            {
                'id': 'final_table',
                'name': '思维发展型课堂教学设计模板（最终表格版）',
                'description': '使用正确表格循环语法的教学设计模板，确保不同行数据在不同行显示',
                'path': 'templates/final_table_template.docx'
            }
        ]
        
        return jsonify({
            'success': True,
            'templates': templates
        })
    except Exception as e:
        logger.error(f"获取模板列表失败: {e}")
        return jsonify({'success': False, 'error': str(e)})

@app.route('/status/<task_id>')
def get_task_status(task_id):
    """获取任务状态"""
    try:
        # 查找对应的文件信息
        file_info = None
        for file_id, info in uploaded_files.items():
            if info.get('task_id') == task_id:
                file_info = info
                break
        
        if not file_info:
            return jsonify({'success': False, 'error': '任务不存在'})
        
        # 检查任务状态
        status = file_info.get('status', 'processing')
        
        if status == 'processing':
            # 如果状态是processing，说明还没有开始AI处理，直接返回处理中状态
            return jsonify({
                'success': True,
                'status': 'processing',
                'progress': 50,
                'message': '文件正在处理中...'
            })
        elif status == 'completed':
            return jsonify({
                'success': True,
                'status': 'completed',
                'progress': 100,
                'result': file_info.get('result', '处理完成')
            })
        else:
            return jsonify({
                'success': True,
                'status': status,
                'progress': 50
            })
        
    except Exception as e:
        logger.error(f"获取任务状态失败: {e}")
        return jsonify({'success': False, 'error': str(e)})

@app.route('/download/<task_id>')
def download_file(task_id):
    """下载生成的文件"""
    try:
        # 查找对应的文件信息
        file_info = None
        for file_id, info in uploaded_files.items():
            if info.get('task_id') == task_id:
                file_info = info
                break
        
        if not file_info:
            return jsonify({'success': False, 'error': '文件不存在'}), 404
        
        # 检查是否有生成的内容
        if not file_info.get('result'):
            return jsonify({'success': False, 'error': '文件尚未生成'}), 400
        
        # 生成Word文档
        from docx import Document
        from docx.shared import Inches
        import tempfile
        import os
        
        # 创建临时文件
        temp_file = tempfile.NamedTemporaryFile(delete=False, suffix='.docx')
        temp_path = temp_file.name
        temp_file.close()
        
        # 添加到清理列表
        temp_files.add(temp_path)
        
        # 创建Word文档
        doc = Document()
        
        # 设置文档标题
        doc.add_heading('教学设计文档', 0)
        
        # 添加生成的内容
        content = file_info['result']
        
        # 将内容按段落分割并添加到文档
        paragraphs = content.split('\n\n')
        for para_text in paragraphs:
            if para_text.strip():
                # 检查是否是标题（以【】包围）
                if para_text.strip().startswith('【') and para_text.strip().endswith('】'):
                    doc.add_heading(para_text.strip()[1:-1], level=1)
                else:
                    doc.add_paragraph(para_text.strip())
        
        # 保存文档
        doc.save(temp_path)
        
        # 返回文件
        response = send_file(
            temp_path,
            as_attachment=True,
            download_name=f"教学设计_{file_info['original_filename']}",
            mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
        )
        
        # 设置响应头，下载完成后删除临时文件
        response.headers['X-Temp-File'] = temp_path
        
        return response
        
    except Exception as e:
        logger.error(f"文件下载失败: {e}")
        return jsonify({'success': False, 'error': str(e)}), 500

@app.route('/favicon.ico')
def favicon():
    """网站图标"""
    return send_file('web/static/favicon.ico', mimetype='image/vnd.microsoft.icon')

@app.route('/health')
def health():
    """健康检查接口"""
    return jsonify({
        'status': 'healthy',
        'timestamp': datetime.now().isoformat(),
        'version': '0.1.0'
    })

@app.errorhandler(404)
def not_found(error):
    """404错误处理"""
    return jsonify({'error': '页面不存在'}), 404

@app.errorhandler(500)
def internal_error(error):
    """500错误处理"""
    logger.error(f"内部服务器错误: {error}")
    return jsonify({'error': '内部服务器错误'}), 500

if __name__ == '__main__':
    # 确保必要的目录存在
    os.makedirs('uploads', exist_ok=True)
    os.makedirs('outputs', exist_ok=True)
    os.makedirs('logs', exist_ok=True)
    
    # 启动应用
    port = int(os.getenv('PORT', 5000))
    debug = os.getenv('FLASK_DEBUG', 'False').lower() == 'true'
    
    logger.info(f"启动AI文档助手，端口: {port}, 调试模式: {debug}")
    
    app.run(
        host='0.0.0.0',
        port=port,
        debug=debug,
        threaded=True
    )