#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
创建使用docxtpl表插入功能的模板
使用docxtpl的RichText和表格循环功能
"""

from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.shared import OxmlElement, qn
from docx.oxml import OxmlElement as oxml_element

def create_docxtpl_table_template():
    """创建使用docxtpl表插入功能的模板"""
    doc = Document()
    
    # 设置页面边距
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)
    
    # 标题
    title = doc.add_heading('思维发展型课堂教学设计', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # 基本信息表格
    info_table = doc.add_table(rows=4, cols=4)
    info_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    info_table.style = 'Table Grid'
    
    # 设置表格列宽
    for row in info_table.rows:
        for i, cell in enumerate(row.cells):
            if i == 0 or i == 2:
                cell.width = Inches(1.2)
            else:
                cell.width = Inches(1.8)
    
    # 填充基本信息
    info_data = [
        ['课例名称', '{{ lesson_name }}', '学段年级', '{{ grade_level }}'],
        ['学科', '{{ subject }}', '教材版本', '{{ textbook_version }}'],
        ['课时', '{{ lesson_period }}', '学校', '{{ teacher_school }}'],
        ['教师', '{{ teacher_name }}', '', '']
    ]
    
    for i, row_data in enumerate(info_data):
        for j, cell_data in enumerate(row_data):
            cell = info_table.cell(i, j)
            if j % 2 == 0:  # 标签列
                cell.text = cell_data
                # 设置标签样式
                for paragraph in cell.paragraphs:
                    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    for run in paragraph.runs:
                        run.font.bold = True
                        run.font.size = Pt(12)
            else:  # 内容列
                cell.text = cell_data
                # 设置内容样式
                for paragraph in cell.paragraphs:
                    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    for run in paragraph.runs:
                        run.font.size = Pt(12)
    
    # 添加空行
    doc.add_paragraph()
    
    # 一、课例概述
    doc.add_heading('一、课例概述', level=1)
    doc.add_paragraph('{{ summary }}')
    
    # 二、内容分析
    doc.add_heading('二、内容分析', level=1)
    doc.add_paragraph('{{ content_analysis }}')
    
    # 三、学习者分析
    doc.add_heading('三、学习者分析', level=1)
    doc.add_paragraph('{{ learner_analysis }}')
    
    # 四、学习目标及重难点
    doc.add_heading('四、学习目标及重难点', level=1)
    doc.add_paragraph('{{ learning_objectives }}')
    
    # 五、教学设计思路
    doc.add_heading('五、教学设计思路', level=1)
    doc.add_paragraph('{{ lesson_structure }}')
    
    # 六、学习活动设计（使用表插入形式）
    doc.add_heading('六、学习活动设计', level=1)
    
    # 创建学习活动表格 - 使用docxtpl的表格循环功能
    # 这里我们创建一个包含循环语法的表格
    activities_table = doc.add_table(rows=2, cols=4)  # 表头 + 循环行
    activities_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    activities_table.style = 'Table Grid'
    
    # 设置表格列宽
    for row in activities_table.rows:
        row.cells[0].width = Inches(1.5)  # 环节名称
        row.cells[1].width = Inches(2.5)  # 教师活动
        row.cells[2].width = Inches(2.5)  # 学生活动
        row.cells[3].width = Inches(1.5)  # 活动意图
    
    # 设置表头
    header_cells = activities_table.rows[0].cells
    headers = ['环节名称', '教师活动', '学生活动', '活动意图']
    for i, header in enumerate(headers):
        header_cells[i].text = header
        # 设置表头样式
        for paragraph in header_cells[i].paragraphs:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in paragraph.runs:
                run.font.bold = True
                run.font.size = Pt(12)
    
    # 设置循环行 - 使用docxtpl的表格循环语法
    loop_row = activities_table.rows[1]
    loop_row.cells[0].text = '{% for activity in learning_activities %}{{ activity.name }}{% endfor %}'
    loop_row.cells[1].text = '{% for activity in learning_activities %}{{ activity.teacher_activity }}{% endfor %}'
    loop_row.cells[2].text = '{% for activity in learning_activities %}{{ activity.student_activity }}{% endfor %}'
    loop_row.cells[3].text = '{% for activity in learning_activities %}{{ activity.activity_intent }}{% endfor %}'
    
    # 七、板书设计
    doc.add_heading('七、板书设计', level=1)
    doc.add_paragraph('{{ blackboard_design }}')
    
    # 八、作业与拓展
    doc.add_heading('八、作业与拓展', level=1)
    doc.add_paragraph('{{ homework_extension }}')
    
    # 九、学习素材设计
    doc.add_heading('九、学习素材设计', level=1)
    doc.add_paragraph('{{ materials_design }}')
    
    # 十、思维训练点设计
    doc.add_heading('十、思维训练点设计', level=1)
    
    # 创建思维训练点表格
    thinking_table = doc.add_table(rows=2, cols=2)
    thinking_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    thinking_table.style = 'Table Grid'
    
    # 设置表格列宽
    for row in thinking_table.rows:
        row.cells[0].width = Inches(2)  # 训练点类型
        row.cells[1].width = Inches(4)  # 描述
    
    # 设置表头
    thinking_header_cells = thinking_table.rows[0].cells
    thinking_headers = ['训练点类型', '具体描述']
    for i, header in enumerate(thinking_headers):
        thinking_header_cells[i].text = header
        # 设置表头样式
        for paragraph in thinking_header_cells[i].paragraphs:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in paragraph.runs:
                run.font.bold = True
                run.font.size = Pt(12)
    
    # 设置循环行
    thinking_loop_row = thinking_table.rows[1]
    thinking_loop_row.cells[0].text = '{% for point in reflection_thinking_points %}{{ point.point_type }}{% endfor %}'
    thinking_loop_row.cells[1].text = '{% for point in reflection_thinking_points %}{{ point.description }}{% endfor %}'
    
    # 保存文档
    output_path = 'templates/docxtpl_table_teaching_design_template.docx'
    doc.save(output_path)
    print(f"docxtpl表插入模板文件已创建: {output_path}")
    
    return output_path

if __name__ == "__main__":
    create_docxtpl_table_template()
