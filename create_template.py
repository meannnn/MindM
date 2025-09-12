#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
创建教学设计Word模板
使用python-docx创建包含占位符的Word模板文件
"""

from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.shared import OxmlElement, qn

def create_teaching_design_template():
    """创建教学设计模板"""
    
    # 创建新文档
    doc = Document()
    
    # 设置文档标题
    title = doc.add_heading('思维发展型课堂教学设计', 0)
    title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    
    # 添加基本信息表格
    info_table = doc.add_table(rows=4, cols=2)
    info_table.style = 'Table Grid'
    info_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    
    # 设置表格内容
    info_data = [
        ['课例名称', '{{lesson_name}}'],
        ['学段年级', '{{grade_level}}'],
        ['学科', '{{subject}}'],
        ['教材版本', '{{textbook_version}}'],
        ['课时说明', '{{lesson_period}}'],
        ['教师单位', '{{teacher_school}}'],
        ['教师姓名', '{{teacher_name}}']
    ]
    
    for i, (label, value) in enumerate(info_data):
        if i < len(info_table.rows):
            row = info_table.rows[i]
            row.cells[0].text = label
            row.cells[1].text = value
            # 设置第一列为粗体
            for paragraph in row.cells[0].paragraphs:
                for run in paragraph.runs:
                    run.bold = True
    
    # 添加空行
    doc.add_paragraph()
    
    # 添加各个部分
    sections = [
        ('【摘要】', '{{summary}}'),
        ('【教学内容分析】', '{{content_analysis}}'),
        ('【学习者分析】', '{{learner_analysis}}'),
        ('【学习目标及重难点】', '{{learning_objectives}}'),
        ('【课例结构】', '{{lesson_structure}}'),
        ('【学习活动设计】', '{{learning_activities}}'),
        ('【板书设计】', '{{blackboard_design}}'),
        ('【作业与拓展学习设计】', '{{homework_extension}}'),
        ('【素材设计】', '{{materials_design}}'),
        ('【反思：思维训练点】', '{{reflection_thinking_points}}')
    ]
    
    for section_title, placeholder in sections:
        # 添加节标题
        heading = doc.add_heading(section_title, level=1)
        
        # 添加占位符内容
        if section_title == '【学习活动设计】':
            # 为学习活动设计创建表格
            activities_table = doc.add_table(rows=3, cols=3)
            activities_table.style = 'Table Grid'
            
            # 设置表头
            header_row = activities_table.rows[0]
            header_row.cells[0].text = '环节'
            header_row.cells[1].text = '教师活动'
            header_row.cells[2].text = '学生活动'
            
            # 设置表头格式
            for cell in header_row.cells:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.bold = True
            
            # 添加示例行
            for i in range(1, 3):
                row = activities_table.rows[i]
                row.cells[0].text = f'{{{{activity_{i}_name}}}}'
                row.cells[1].text = f'{{{{activity_{i}_teacher}}}}'
                row.cells[2].text = f'{{{{activity_{i}_student}}}}'
            
            # 添加活动意图说明
            doc.add_paragraph('活动意图说明：{{activity_intent}}')
            
        elif section_title == '【反思：思维训练点】':
            # 为思维训练点创建表格
            thinking_table = doc.add_table(rows=4, cols=2)
            thinking_table.style = 'Table Grid'
            
            # 设置表头
            header_row = thinking_table.rows[0]
            header_row.cells[0].text = '思维训练点'
            header_row.cells[1].text = '说明'
            
            # 设置表头格式
            for cell in header_row.cells:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.bold = True
            
            # 添加思维训练点
            thinking_points = [
                ('认知冲突', '{{cognitive_conflict}}'),
                ('思维图示', '{{thinking_diagram}}'),
                ('变式运用', '{{variation_application}}')
            ]
            
            for i, (point, description) in enumerate(thinking_points, 1):
                row = thinking_table.rows[i]
                row.cells[0].text = point
                row.cells[1].text = description
        
        else:
            # 普通段落
            doc.add_paragraph(placeholder)
        
        # 添加空行
        doc.add_paragraph()
    
    # 保存文档
    output_path = 'templates/teaching_design_template.docx'
    doc.save(output_path)
    print(f"模板文件已创建: {output_path}")

if __name__ == "__main__":
    create_teaching_design_template()
