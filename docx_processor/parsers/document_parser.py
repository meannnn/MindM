"""
Word文档解析器

基于python-docx库实现Word文档的全面解析功能。
参考：https://python-docx.readthedocs.io/en/latest/

功能包括：
- 提取文档基本信息
- 解析段落和文本内容
- 提取表格数据
- 识别图片和形状
- 分析文档结构
"""

import os
import logging
from typing import Dict, List, Any, Optional, Tuple
from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

logger = logging.getLogger(__name__)


class DocumentParser:
    """Word文档解析器"""
    
    def __init__(self, file_path: str):
        """
        初始化文档解析器
        
        Args:
            file_path: Word文档文件路径
        """
        self.file_path = file_path
        self.document = None
        self.parsed_data = {}
        
        if not os.path.exists(file_path):
            raise FileNotFoundError(f"文件不存在: {file_path}")
            
        if not file_path.lower().endswith('.docx'):
            raise ValueError("只支持.docx格式的Word文档")
    
    def load_document(self) -> bool:
        """
        加载Word文档
        
        Returns:
            bool: 加载是否成功
        """
        try:
            self.document = Document(self.file_path)
            logger.info(f"成功加载文档: {self.file_path}")
            return True
        except Exception as e:
            logger.error(f"加载文档失败: {e}")
            return False
    
    def parse_document(self) -> Dict[str, Any]:
        """
        解析整个文档
        
        Returns:
            Dict[str, Any]: 解析后的文档数据
        """
        if not self.document:
            if not self.load_document():
                return {}
        
        try:
            self.parsed_data = {
                'basic_info': self._extract_basic_info(),
                'paragraphs': self._extract_paragraphs(),
                'tables': self._extract_tables(),
                'images': self._extract_images(),
                'sections': self._extract_sections(),
                'styles': self._extract_styles(),
                'structure': self._analyze_structure()
            }
            
            logger.info("文档解析完成")
            return self.parsed_data
            
        except Exception as e:
            logger.error(f"文档解析失败: {e}")
            return {}
    
    def _extract_basic_info(self) -> Dict[str, Any]:
        """提取文档基本信息"""
        try:
            core_props = self.document.core_properties
            
            return {
                'title': core_props.title or '',
                'author': core_props.author or '',
                'subject': core_props.subject or '',
                'keywords': core_props.keywords or '',
                'comments': core_props.comments or '',
                'created': core_props.created,
                'modified': core_props.modified,
                'last_modified_by': core_props.last_modified_by or '',
                'revision': core_props.revision or 0,
                'version': core_props.version or '',
                'file_size': os.path.getsize(self.file_path),
                'paragraph_count': len(self.document.paragraphs),
                'table_count': len(self.document.tables)
            }
        except Exception as e:
            logger.error(f"提取基本信息失败: {e}")
            return {}
    
    def _extract_paragraphs(self) -> List[Dict[str, Any]]:
        """提取段落信息"""
        paragraphs_data = []
        
        try:
            for i, paragraph in enumerate(self.document.paragraphs):
                if paragraph.text.strip():  # 只处理非空段落
                    para_data = {
                        'index': i,
                        'text': paragraph.text,
                        'style': paragraph.style.name if paragraph.style else 'Normal',
                        'alignment': self._get_alignment_name(paragraph.alignment),
                        'runs': self._extract_runs(paragraph),
                        'level': getattr(paragraph, 'level', 0),
                        'is_heading': paragraph.style.name.startswith('Heading') if paragraph.style else False
                    }
                    paragraphs_data.append(para_data)
            
            logger.info(f"提取了 {len(paragraphs_data)} 个段落")
            return paragraphs_data
            
        except Exception as e:
            logger.error(f"提取段落信息失败: {e}")
            return []
    
    def _extract_runs(self, paragraph) -> List[Dict[str, Any]]:
        """提取段落中的文本运行信息"""
        runs_data = []
        
        try:
            for run in paragraph.runs:
                if run.text.strip():
                    run_data = {
                        'text': run.text,
                        'bold': run.bold,
                        'italic': run.italic,
                        'underline': run.underline,
                        'font_name': run.font.name,
                        'font_size': run.font.size.pt if run.font.size else None,
                        'font_color': str(run.font.color.rgb) if run.font.color and run.font.color.rgb else None
                    }
                    runs_data.append(run_data)
            
            return runs_data
            
        except Exception as e:
            logger.error(f"提取文本运行信息失败: {e}")
            return []
    
    def _extract_tables(self) -> List[Dict[str, Any]]:
        """提取表格信息"""
        tables_data = []
        
        try:
            for i, table in enumerate(self.document.tables):
                table_data = {
                    'index': i,
                    'rows': len(table.rows),
                    'columns': len(table.columns),
                    'data': self._extract_table_data(table),
                    'style': table.style.name if table.style else 'Table Grid',
                    'alignment': self._get_table_alignment_name(table.alignment)
                }
                tables_data.append(table_data)
            
            logger.info(f"提取了 {len(tables_data)} 个表格")
            return tables_data
            
        except Exception as e:
            logger.error(f"提取表格信息失败: {e}")
            return []
    
    def _extract_table_data(self, table) -> List[List[str]]:
        """提取表格数据"""
        table_data = []
        
        try:
            for row in table.rows:
                row_data = []
                for cell in row.cells:
                    cell_text = cell.text.strip()
                    row_data.append(cell_text)
                table_data.append(row_data)
            
            return table_data
            
        except Exception as e:
            logger.error(f"提取表格数据失败: {e}")
            return []
    
    def _extract_images(self) -> List[Dict[str, Any]]:
        """提取图片信息"""
        images_data = []
        
        try:
            # 提取内联形状（图片）
            for i, shape in enumerate(self.document.inline_shapes):
                if hasattr(shape, 'image'):
                    image_data = {
                        'index': i,
                        'type': 'inline',
                        'width': shape.width,
                        'height': shape.height,
                        'filename': getattr(shape.image, 'filename', 'unknown')
                    }
                    images_data.append(image_data)
            
            # 提取段落中的图片
            for para_idx, paragraph in enumerate(self.document.paragraphs):
                for run in paragraph.runs:
                    if run._element.xpath('.//a:blip'):
                        image_data = {
                            'index': len(images_data),
                            'type': 'paragraph',
                            'paragraph_index': para_idx,
                            'run_text': run.text
                        }
                        images_data.append(image_data)
            
            logger.info(f"提取了 {len(images_data)} 个图片")
            return images_data
            
        except Exception as e:
            logger.error(f"提取图片信息失败: {e}")
            return []
    
    def _extract_sections(self) -> List[Dict[str, Any]]:
        """提取章节信息"""
        sections_data = []
        
        try:
            for i, section in enumerate(self.document.sections):
                section_data = {
                    'index': i,
                    'page_width': section.page_width,
                    'page_height': section.page_height,
                    'left_margin': section.left_margin,
                    'right_margin': section.right_margin,
                    'top_margin': section.top_margin,
                    'bottom_margin': section.bottom_margin,
                    'orientation': 'landscape' if section.orientation == 1 else 'portrait'
                }
                sections_data.append(section_data)
            
            logger.info(f"提取了 {len(sections_data)} 个章节")
            return sections_data
            
        except Exception as e:
            logger.error(f"提取章节信息失败: {e}")
            return []
    
    def _extract_styles(self) -> Dict[str, Any]:
        """提取样式信息"""
        styles_data = {
            'paragraph_styles': [],
            'character_styles': [],
            'table_styles': []
        }
        
        try:
            for style in self.document.styles:
                style_data = {
                    'name': style.name,
                    'type': str(style.type),
                    'builtin': style.builtin,
                    'hidden': style.hidden
                }
                
                if style.type == 1:  # 段落样式
                    styles_data['paragraph_styles'].append(style_data)
                elif style.type == 2:  # 字符样式
                    styles_data['character_styles'].append(style_data)
                elif style.type == 3:  # 表格样式
                    styles_data['table_styles'].append(style_data)
            
            logger.info(f"提取了样式信息: 段落样式 {len(styles_data['paragraph_styles'])} 个, "
                       f"字符样式 {len(styles_data['character_styles'])} 个, "
                       f"表格样式 {len(styles_data['table_styles'])} 个")
            
            return styles_data
            
        except Exception as e:
            logger.error(f"提取样式信息失败: {e}")
            return styles_data
    
    def _analyze_structure(self) -> Dict[str, Any]:
        """分析文档结构"""
        structure_data = {
            'headings': [],
            'lists': [],
            'tables': [],
            'images': [],
            'page_breaks': []
        }
        
        try:
            # 分析标题结构
            for i, paragraph in enumerate(self.document.paragraphs):
                if paragraph.style and paragraph.style.name.startswith('Heading'):
                    heading_data = {
                        'index': i,
                        'text': paragraph.text,
                        'level': int(paragraph.style.name.split()[-1]) if paragraph.style.name.split()[-1].isdigit() else 1,
                        'style': paragraph.style.name
                    }
                    structure_data['headings'].append(heading_data)
            
            # 分析列表
            for i, paragraph in enumerate(self.document.paragraphs):
                if paragraph.style and ('List' in paragraph.style.name or 'Bullet' in paragraph.style.name):
                    list_data = {
                        'index': i,
                        'text': paragraph.text,
                        'style': paragraph.style.name
                    }
                    structure_data['lists'].append(list_data)
            
            # 分析表格位置
            for i, table in enumerate(self.document.tables):
                # 找到表格前的段落来确定位置
                table_data = {
                    'index': i,
                    'rows': len(table.rows),
                    'columns': len(table.columns)
                }
                structure_data['tables'].append(table_data)
            
            logger.info(f"文档结构分析完成: 标题 {len(structure_data['headings'])} 个, "
                       f"列表 {len(structure_data['lists'])} 个, "
                       f"表格 {len(structure_data['tables'])} 个")
            
            return structure_data
            
        except Exception as e:
            logger.error(f"分析文档结构失败: {e}")
            return structure_data
    
    def _get_alignment_name(self, alignment) -> str:
        """获取对齐方式名称"""
        if alignment == WD_PARAGRAPH_ALIGNMENT.LEFT:
            return 'left'
        elif alignment == WD_PARAGRAPH_ALIGNMENT.CENTER:
            return 'center'
        elif alignment == WD_PARAGRAPH_ALIGNMENT.RIGHT:
            return 'right'
        elif alignment == WD_PARAGRAPH_ALIGNMENT.JUSTIFY:
            return 'justify'
        else:
            return 'left'
    
    def _get_table_alignment_name(self, alignment) -> str:
        """获取表格对齐方式名称"""
        if alignment == WD_TABLE_ALIGNMENT.LEFT:
            return 'left'
        elif alignment == WD_TABLE_ALIGNMENT.CENTER:
            return 'center'
        elif alignment == WD_TABLE_ALIGNMENT.RIGHT:
            return 'right'
        else:
            return 'left'
    
    def get_text_content(self) -> str:
        """获取纯文本内容"""
        if not self.parsed_data:
            self.parse_document()
        
        text_content = []
        for paragraph in self.parsed_data.get('paragraphs', []):
            text_content.append(paragraph['text'])
        
        return '\n'.join(text_content)
    
    def get_structured_content(self) -> Dict[str, Any]:
        """获取结构化内容，用于AI分析"""
        if not self.parsed_data:
            self.parse_document()
        
        structured_content = {
            'title': self.parsed_data.get('basic_info', {}).get('title', ''),
            'author': self.parsed_data.get('basic_info', {}).get('author', ''),
            'text_content': self.get_text_content(),
            'headings': [h['text'] for h in self.parsed_data.get('structure', {}).get('headings', [])],
            'tables': self.parsed_data.get('tables', []),
            'lists': [l['text'] for l in self.parsed_data.get('structure', {}).get('lists', [])]
        }
        
        return structured_content
    
    def save_parsed_data(self, output_path: str) -> bool:
        """保存解析数据到JSON文件"""
        import json
        
        try:
            with open(output_path, 'w', encoding='utf-8') as f:
                json.dump(self.parsed_data, f, ensure_ascii=False, indent=2, default=str)
            
            logger.info(f"解析数据已保存到: {output_path}")
            return True
            
        except Exception as e:
            logger.error(f"保存解析数据失败: {e}")
            return False


# 使用示例
if __name__ == "__main__":
    # 配置日志
    logging.basicConfig(level=logging.INFO)
    
    # 创建解析器实例
    parser = DocumentParser("example.docx")
    
    # 解析文档
    data = parser.parse_document()
    
    # 打印基本信息
    basic_info = data.get('basic_info', {})
    print(f"文档标题: {basic_info.get('title', '无标题')}")
    print(f"作者: {basic_info.get('author', '未知')}")
    print(f"段落数: {basic_info.get('paragraph_count', 0)}")
    print(f"表格数: {basic_info.get('table_count', 0)}")
    
    # 获取结构化内容
    structured = parser.get_structured_content()
    print(f"文本内容长度: {len(structured['text_content'])}")
    print(f"标题数量: {len(structured['headings'])}")