"""
格式解析器

专门用于解析Word文档中的格式信息，包括字体、段落格式、样式等。
"""

import logging
from typing import Dict, List, Any, Optional
from docx import Document
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT, WD_UNDERLINE
from docx.enum.style import WD_STYLE_TYPE
from docx.shared import RGBColor
from docx.oxml.ns import qn

logger = logging.getLogger(__name__)


class FormatParser:
    """格式解析器"""
    
    def __init__(self, document: Document):
        """
        初始化格式解析器
        
        Args:
            document: Word文档对象
        """
        self.document = document
    
    def parse_all_formats(self) -> Dict[str, Any]:
        """
        解析所有格式信息
        
        Returns:
            Dict[str, Any]: 所有格式信息
        """
        try:
            format_data = {
                'styles': self._parse_styles(),
                'fonts': self._parse_fonts(),
                'paragraph_formats': self._parse_paragraph_formats(),
                'character_formats': self._parse_character_formats(),
                'table_formats': self._parse_table_formats(),
                'section_formats': self._parse_section_formats(),
                'theme': self._parse_theme()
            }
            
            logger.info("格式解析完成")
            return format_data
            
        except Exception as e:
            logger.error(f"格式解析失败: {e}")
            return {}
    
    def _parse_styles(self) -> Dict[str, Any]:
        """
        解析样式信息
        
        Returns:
            Dict[str, Any]: 样式信息
        """
        styles_data = {
            'paragraph_styles': [],
            'character_styles': [],
            'table_styles': [],
            'list_styles': []
        }
        
        try:
            for style in self.document.styles:
                style_info = {
                    'name': style.name,
                    'type': self._get_style_type_name(style.type),
                    'builtin': style.builtin,
                    'hidden': style.hidden,
                    'locked': getattr(style, 'locked', False),
                    'priority': getattr(style, 'priority', 0),
                    'format': self._get_style_format(style)
                }
                
                if style.type == WD_STYLE_TYPE.PARAGRAPH:
                    styles_data['paragraph_styles'].append(style_info)
                elif style.type == WD_STYLE_TYPE.CHARACTER:
                    styles_data['character_styles'].append(style_info)
                elif style.type == WD_STYLE_TYPE.TABLE:
                    styles_data['table_styles'].append(style_info)
                elif style.type == WD_STYLE_TYPE.LIST:
                    styles_data['list_styles'].append(style_info)
            
            logger.info(f"解析了 {len(self.document.styles)} 个样式")
            return styles_data
            
        except Exception as e:
            logger.error(f"解析样式失败: {e}")
            return styles_data
    
    def _get_style_type_name(self, style_type) -> str:
        """获取样式类型名称"""
        type_names = {
            WD_STYLE_TYPE.PARAGRAPH: 'paragraph',
            WD_STYLE_TYPE.CHARACTER: 'character',
            WD_STYLE_TYPE.TABLE: 'table',
            WD_STYLE_TYPE.LIST: 'list'
        }
        return type_names.get(style_type, 'unknown')
    
    def _get_style_format(self, style) -> Dict[str, Any]:
        """
        获取样式格式信息
        
        Args:
            style: 样式对象
            
        Returns:
            Dict[str, Any]: 格式信息
        """
        try:
            format_info = {}
            
            # 段落格式
            if hasattr(style, 'paragraph_format'):
                para_format = style.paragraph_format
                format_info['paragraph'] = {
                    'alignment': self._get_alignment_name(para_format.alignment),
                    'space_before': str(para_format.space_before) if para_format.space_before else None,
                    'space_after': str(para_format.space_after) if para_format.space_after else None,
                    'line_spacing': str(para_format.line_spacing) if para_format.line_spacing else None,
                    'first_line_indent': str(para_format.first_line_indent) if para_format.first_line_indent else None,
                    'left_indent': str(para_format.left_indent) if para_format.left_indent else None,
                    'right_indent': str(para_format.right_indent) if para_format.right_indent else None
                }
            
            # 字体格式
            if hasattr(style, 'font'):
                font = style.font
                format_info['font'] = {
                    'name': font.name,
                    'size': str(font.size) if font.size else None,
                    'bold': font.bold,
                    'italic': font.italic,
                    'underline': self._get_underline_name(font.underline),
                    'color': str(font.color.rgb) if font.color and font.color.rgb else None
                }
            
            return format_info
            
        except Exception as e:
            logger.error(f"获取样式格式失败: {e}")
            return {}
    
    def _parse_fonts(self) -> Dict[str, Any]:
        """
        解析字体信息
        
        Returns:
            Dict[str, Any]: 字体信息
        """
        fonts_data = {
            'used_fonts': set(),
            'font_sizes': set(),
            'font_colors': set(),
            'font_styles': {
                'bold': 0,
                'italic': 0,
                'underline': 0
            }
        }
        
        try:
            # 遍历所有段落
            for paragraph in self.document.paragraphs:
                for run in paragraph.runs:
                    if run.text.strip():
                        # 字体名称
                        if run.font.name:
                            fonts_data['used_fonts'].add(run.font.name)
                        
                        # 字体大小
                        if run.font.size:
                            fonts_data['font_sizes'].add(str(run.font.size))
                        
                        # 字体颜色
                        if run.font.color and run.font.color.rgb:
                            fonts_data['font_colors'].add(str(run.font.color.rgb))
                        
                        # 字体样式
                        if run.bold:
                            fonts_data['font_styles']['bold'] += 1
                        if run.italic:
                            fonts_data['font_styles']['italic'] += 1
                        if run.underline:
                            fonts_data['font_styles']['underline'] += 1
            
            # 转换为列表
            fonts_data['used_fonts'] = list(fonts_data['used_fonts'])
            fonts_data['font_sizes'] = list(fonts_data['font_sizes'])
            fonts_data['font_colors'] = list(fonts_data['font_colors'])
            
            logger.info(f"解析了 {len(fonts_data['used_fonts'])} 种字体")
            return fonts_data
            
        except Exception as e:
            logger.error(f"解析字体失败: {e}")
            return fonts_data
    
    def _parse_paragraph_formats(self) -> Dict[str, Any]:
        """
        解析段落格式信息
        
        Returns:
            Dict[str, Any]: 段落格式信息
        """
        para_formats = {
            'alignments': {},
            'indents': {
                'left': [],
                'right': [],
                'first_line': []
            },
            'spacing': {
                'before': [],
                'after': [],
                'line': []
            }
        }
        
        try:
            for paragraph in self.document.paragraphs:
                if paragraph.text.strip():
                    para_format = paragraph.paragraph_format
                    
                    # 对齐方式
                    alignment = self._get_alignment_name(para_format.alignment)
                    para_formats['alignments'][alignment] = para_formats['alignments'].get(alignment, 0) + 1
                    
                    # 缩进
                    if para_format.left_indent:
                        para_formats['indents']['left'].append(str(para_format.left_indent))
                    if para_format.right_indent:
                        para_formats['indents']['right'].append(str(para_format.right_indent))
                    if para_format.first_line_indent:
                        para_formats['indents']['first_line'].append(str(para_format.first_line_indent))
                    
                    # 间距
                    if para_format.space_before:
                        para_formats['spacing']['before'].append(str(para_format.space_before))
                    if para_format.space_after:
                        para_formats['spacing']['after'].append(str(para_format.space_after))
                    if para_format.line_spacing:
                        para_formats['spacing']['line'].append(str(para_format.line_spacing))
            
            # 去重
            for indent_type in para_formats['indents']:
                para_formats['indents'][indent_type] = list(set(para_formats['indents'][indent_type]))
            
            for spacing_type in para_formats['spacing']:
                para_formats['spacing'][spacing_type] = list(set(para_formats['spacing'][spacing_type]))
            
            logger.info("段落格式解析完成")
            return para_formats
            
        except Exception as e:
            logger.error(f"解析段落格式失败: {e}")
            return para_formats
    
    def _parse_character_formats(self) -> Dict[str, Any]:
        """
        解析字符格式信息
        
        Returns:
            Dict[str, Any]: 字符格式信息
        """
        char_formats = {
            'bold_count': 0,
            'italic_count': 0,
            'underline_count': 0,
            'underline_types': {},
            'color_usage': {},
            'size_usage': {}
        }
        
        try:
            for paragraph in self.document.paragraphs:
                for run in paragraph.runs:
                    if run.text.strip():
                        # 粗体
                        if run.bold:
                            char_formats['bold_count'] += 1
                        
                        # 斜体
                        if run.italic:
                            char_formats['italic_count'] += 1
                        
                        # 下划线
                        if run.underline:
                            char_formats['underline_count'] += 1
                            underline_type = self._get_underline_name(run.underline)
                            char_formats['underline_types'][underline_type] = char_formats['underline_types'].get(underline_type, 0) + 1
                        
                        # 颜色使用
                        if run.font.color and run.font.color.rgb:
                            color = str(run.font.color.rgb)
                            char_formats['color_usage'][color] = char_formats['color_usage'].get(color, 0) + 1
                        
                        # 大小使用
                        if run.font.size:
                            size = str(run.font.size)
                            char_formats['size_usage'][size] = char_formats['size_usage'].get(size, 0) + 1
            
            logger.info("字符格式解析完成")
            return char_formats
            
        except Exception as e:
            logger.error(f"解析字符格式失败: {e}")
            return char_formats
    
    def _parse_table_formats(self) -> Dict[str, Any]:
        """
        解析表格格式信息
        
        Returns:
            Dict[str, Any]: 表格格式信息
        """
        table_formats = {
            'styles': {},
            'alignments': {},
            'borders': {},
            'shading': {}
        }
        
        try:
            for table in self.document.tables:
                # 表格样式
                style_name = table.style.name if table.style else 'Table Grid'
                table_formats['styles'][style_name] = table_formats['styles'].get(style_name, 0) + 1
                
                # 表格对齐
                alignment = self._get_table_alignment_name(table.alignment)
                table_formats['alignments'][alignment] = table_formats['alignments'].get(alignment, 0) + 1
                
                # 分析单元格格式
                for row in table.rows:
                    for cell in row.cells:
                        # 单元格底纹
                        shading = self._get_cell_shading(cell)
                        if shading:
                            shading_key = f"{shading.get('color', 'none')}_{shading.get('fill', 'none')}"
                            table_formats['shading'][shading_key] = table_formats['shading'].get(shading_key, 0) + 1
            
            logger.info("表格格式解析完成")
            return table_formats
            
        except Exception as e:
            logger.error(f"解析表格格式失败: {e}")
            return table_formats
    
    def _parse_section_formats(self) -> Dict[str, Any]:
        """
        解析章节格式信息
        
        Returns:
            Dict[str, Any]: 章节格式信息
        """
        section_formats = {
            'page_sizes': [],
            'margins': {
                'left': [],
                'right': [],
                'top': [],
                'bottom': []
            },
            'orientations': {},
            'headers_footers': {
                'has_header': 0,
                'has_footer': 0
            }
        }
        
        try:
            for section in self.document.sections:
                # 页面大小
                page_size = f"{section.page_width} x {section.page_height}"
                section_formats['page_sizes'].append(page_size)
                
                # 页边距
                section_formats['margins']['left'].append(str(section.left_margin))
                section_formats['margins']['right'].append(str(section.right_margin))
                section_formats['margins']['top'].append(str(section.top_margin))
                section_formats['margins']['bottom'].append(str(section.bottom_margin))
                
                # 页面方向
                orientation = 'landscape' if section.orientation == 1 else 'portrait'
                section_formats['orientations'][orientation] = section_formats['orientations'].get(orientation, 0) + 1
                
                # 页眉页脚
                if section.header:
                    section_formats['headers_footers']['has_header'] += 1
                if section.footer:
                    section_formats['headers_footers']['has_footer'] += 1
            
            # 去重
            section_formats['page_sizes'] = list(set(section_formats['page_sizes']))
            for margin_type in section_formats['margins']:
                section_formats['margins'][margin_type] = list(set(section_formats['margins'][margin_type]))
            
            logger.info("章节格式解析完成")
            return section_formats
            
        except Exception as e:
            logger.error(f"解析章节格式失败: {e}")
            return section_formats
    
    def _parse_theme(self) -> Dict[str, Any]:
        """
        解析主题信息
        
        Returns:
            Dict[str, Any]: 主题信息
        """
        theme_data = {
            'colors': {},
            'fonts': {},
            'effects': {}
        }
        
        try:
            # 注意：python-docx对主题的支持有限
            # 这里主要提取一些基本的主题相关信息
            
            # 从样式中提取主题颜色
            for style in self.document.styles:
                if hasattr(style, 'font') and style.font.color and style.font.color.rgb:
                    color = str(style.font.color.rgb)
                    theme_data['colors'][color] = theme_data['colors'].get(color, 0) + 1
            
            # 从样式中提取主题字体
            for style in self.document.styles:
                if hasattr(style, 'font') and style.font.name:
                    font = style.font.name
                    theme_data['fonts'][font] = theme_data['fonts'].get(font, 0) + 1
            
            logger.info("主题解析完成")
            return theme_data
            
        except Exception as e:
            logger.error(f"解析主题失败: {e}")
            return theme_data
    
    def _get_alignment_name(self, alignment) -> str:
        """获取对齐方式名称"""
        if alignment == WD_PARAGRAPH_ALIGNMENT.LEFT:
            return 'left'
        elif alignment == WD_PARAGRAPH_ALIGNMENT.CENTER:
            return 'center'
        elif alignment == WD_PARAGRAPH_ALIGNMENT.RIGHT:
            return 'right'
        elif alignment == WD_PARAGRAPH_ALIGNMENT.JUSTIFY:
            return 'justify'
        else:
            return 'left'
    
    def _get_underline_name(self, underline) -> str:
        """获取下划线类型名称"""
        if underline == WD_UNDERLINE.SINGLE:
            return 'single'
        elif underline == WD_UNDERLINE.DOUBLE:
            return 'double'
        elif underline == WD_UNDERLINE.THICK:
            return 'thick'
        elif underline == WD_UNDERLINE.DOTTED:
            return 'dotted'
        elif underline == WD_UNDERLINE.DASHED:
            return 'dashed'
        else:
            return 'none'
    
    def _get_table_alignment_name(self, alignment) -> str:
        """获取表格对齐方式名称"""
        from docx.enum.table import WD_TABLE_ALIGNMENT
        
        if alignment == WD_TABLE_ALIGNMENT.LEFT:
            return 'left'
        elif alignment == WD_TABLE_ALIGNMENT.CENTER:
            return 'center'
        elif alignment == WD_TABLE_ALIGNMENT.RIGHT:
            return 'right'
        else:
            return 'left'
    
    def _get_cell_shading(self, cell) -> Dict[str, Any]:
        """获取单元格底纹信息"""
        try:
            shading = cell._tc.get_or_add_tcPr().get_or_add_shd()
            return {
                'color': str(shading.val) if shading.val else None,
                'fill': str(shading.fill) if shading.fill else None
            }
        except:
            return {}
    
    def get_format_summary(self) -> Dict[str, Any]:
        """
        获取格式汇总信息
        
        Returns:
            Dict[str, Any]: 格式汇总信息
        """
        try:
            format_data = self.parse_all_formats()
            
            summary = {
                'total_styles': len(format_data.get('styles', {}).get('paragraph_styles', [])) + 
                               len(format_data.get('styles', {}).get('character_styles', [])) +
                               len(format_data.get('styles', {}).get('table_styles', [])),
                'font_count': len(format_data.get('fonts', {}).get('used_fonts', [])),
                'color_count': len(format_data.get('fonts', {}).get('font_colors', [])),
                'size_count': len(format_data.get('fonts', {}).get('font_sizes', [])),
                'table_count': len(self.document.tables),
                'section_count': len(self.document.sections)
            }
            
            return summary
            
        except Exception as e:
            logger.error(f"获取格式汇总失败: {e}")
            return {}
    
    def export_formats_to_dict(self) -> Dict[str, Any]:
        """
        导出格式信息为字典格式
        
        Returns:
            Dict[str, Any]: 格式信息字典
        """
        try:
            format_data = self.parse_all_formats()
            
            # 转换为更简洁的格式
            export_data = {
                'styles': {
                    'paragraph': [s['name'] for s in format_data.get('styles', {}).get('paragraph_styles', [])],
                    'character': [s['name'] for s in format_data.get('styles', {}).get('character_styles', [])],
                    'table': [s['name'] for s in format_data.get('styles', {}).get('table_styles', [])]
                },
                'fonts': {
                    'names': format_data.get('fonts', {}).get('used_fonts', []),
                    'sizes': format_data.get('fonts', {}).get('font_sizes', []),
                    'colors': format_data.get('fonts', {}).get('font_colors', [])
                },
                'paragraphs': {
                    'alignments': format_data.get('paragraph_formats', {}).get('alignments', {}),
                    'indents': format_data.get('paragraph_formats', {}).get('indents', {}),
                    'spacing': format_data.get('paragraph_formats', {}).get('spacing', {})
                },
                'characters': {
                    'bold_count': format_data.get('character_formats', {}).get('bold_count', 0),
                    'italic_count': format_data.get('character_formats', {}).get('italic_count', 0),
                    'underline_count': format_data.get('character_formats', {}).get('underline_count', 0)
                }
            }
            
            return export_data
            
        except Exception as e:
            logger.error(f"导出格式信息失败: {e}")
            return {}


# 使用示例
if __name__ == "__main__":
    # 配置日志
    logging.basicConfig(level=logging.INFO)
    
    # 加载文档
    doc = Document("example.docx")
    
    # 创建格式解析器
    parser = FormatParser(doc)
    
    # 解析所有格式
    format_data = parser.parse_all_formats()
    
    # 打印格式信息
    print("样式信息:")
    styles = format_data.get('styles', {})
    print(f"  段落样式: {len(styles.get('paragraph_styles', []))} 个")
    print(f"  字符样式: {len(styles.get('character_styles', []))} 个")
    print(f"  表格样式: {len(styles.get('table_styles', []))} 个")
    
    print("\n字体信息:")
    fonts = format_data.get('fonts', {})
    print(f"  使用字体: {fonts.get('used_fonts', [])}")
    print(f"  字体大小: {fonts.get('font_sizes', [])}")
    print(f"  字体颜色: {fonts.get('font_colors', [])}")
    
    # 获取汇总信息
    summary = parser.get_format_summary()
    print(f"\n格式汇总:")
    print(f"  总样式数: {summary.get('total_styles', 0)}")
    print(f"  字体种类: {summary.get('font_count', 0)}")
    print(f"  颜色种类: {summary.get('color_count', 0)}")
    print(f"  表格数量: {summary.get('table_count', 0)}")